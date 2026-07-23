"""Извлечение текста из DOCX, XLSX и старого XLS.

.docx и .xlsx — это zip-архивы. Текст достаём напрямую (быстро, без OCR).
.xls (старый бинарный формат Excel 97-2003) — через xlrd 1.2.0.

Таблицы спецификации (наименование/характеристики/ед.изм./количество)
парсятся НАПРЯМУЮ по структуре (spec_table.py), минуя LLM — это не только
быстрее и дешевле, но и убирает риск, что модель откажется анализировать
текст из-за формулировок вроде "раствор кислот и щелочей 40%" (совершенно
обычный язык описания товара, но иногда ложно триггерит фильтры безопасности
LLM). Найденная таблица заменяется в тексте короткой пометкой — модель
всё равно знает, что позиции есть, но не видит "опасные" формулировки.

Картинки, вложенные в docx/xlsx, прогоняются через Yandex Vision OCR.
Декоративные элементы (логотипы, печати) обычно очень маленькие по размеру
файла — пропускаем такие, чтобы не тратить OCR-вызовы и не рисковать
зависанием на "битом" изображении. Одна упавшая картинка не роняет весь разбор.

ВАЖНО: старый .doc (Word 97-2003) НЕ поддерживается.
"""
import zipfile
import io

from docx import Document
from openpyxl import load_workbook
import xlrd

from ocr import recognize_image
from spec_table import extract_from_docx_tables, extract_from_xlsx_sheet

IMAGE_EXTS = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".bmp": "image/bmp", ".tiff": "image/tiff"}
MIN_IMAGE_BYTES = 3000  # декоративные логотипы/иконки обычно меньше — реального текста в них нет


def _extract_embedded_images(file_bytes: bytes, media_prefix: str) -> list[tuple[bytes, str]]:
    images = []
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        for name in z.namelist():
            lower = name.lower()
            if not name.startswith(media_prefix):
                continue
            ext = next((e for e in IMAGE_EXTS if lower.endswith(e)), None)
            if not ext:
                continue
            info = z.getinfo(name)
            if info.file_size < MIN_IMAGE_BYTES:
                continue
            images.append((z.read(name), IMAGE_EXTS[ext]))
    return images


async def _ocr_images_text(images: list[tuple[bytes, str]]) -> str:
    parts = []
    for img_bytes, mime_type in images:
        try:
            text = await recognize_image(img_bytes, mime_type)
        except Exception:  # noqa: BLE001
            text = ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


async def extract_docx(file_bytes: bytes) -> tuple[str, list[dict]]:
    """Возвращает (текст для LLM, позиции спецификации если найдены структурно)."""
    doc = Document(io.BytesIO(file_bytes))

    spec_items, spec_tables = extract_from_docx_tables(doc)

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for t_idx, table in enumerate(doc.tables):
        if t_idx in spec_tables:
            parts.append(f"[Таблица спецификации — {len(spec_items)} позиций, обработана отдельно]")
            continue
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))

    text = "\n".join(parts)

    images = _extract_embedded_images(file_bytes, "word/media/")
    ocr_text = await _ocr_images_text(images)
    if ocr_text:
        text += "\n\n[Текст с вложенного изображения]\n" + ocr_text

    return text, spec_items


async def extract_xlsx(file_bytes: bytes) -> tuple[str, list[dict]]:
    """Возвращает (текст для LLM, позиции спецификации если найдены структурно)."""
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)

    parts = []
    spec_items: list[dict] = []
    for sheet in wb.worksheets:
        sheet_spec = extract_from_xlsx_sheet(sheet)
        if sheet_spec:
            spec_items.extend(sheet_spec)
            parts.append(f"--- Лист: {sheet.title} --- [таблица спецификации — {len(sheet_spec)} позиций, обработана отдельно]")
            continue
        parts.append(f"--- Лист: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(parts)

    images = _extract_embedded_images(file_bytes, "xl/media/")
    ocr_text = await _ocr_images_text(images)
    if ocr_text:
        text += "\n\n[Текст с вложенного изображения]\n" + ocr_text

    return text, spec_items


async def extract_xls(file_bytes: bytes) -> tuple[str, list[dict]]:
    """Старый бинарный формат Excel 97-2003. Структурный поиск спецификации не
    делаем (xlrd отдаёт данные иначе) — .xls на площадках закупок это обычно
    расчёт НМЦК, а не спецификация с характеристиками, риск отказа модели ниже."""
    wb = xlrd.open_workbook(file_contents=file_bytes)

    parts = []
    for sheet in wb.sheets():
        parts.append(f"--- Лист: {sheet.name} ---")
        for row_idx in range(sheet.nrows):
            cells = [str(c.value) if c.value != "" else "" for c in sheet.row(row_idx)]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts), []
